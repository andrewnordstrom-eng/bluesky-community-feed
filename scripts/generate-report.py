#!/usr/bin/env python3
"""
Feed Quality Analysis Report Generator

Generates a styled .docx report with charts and tables from the feed's
scoring pipeline data. Pulls data from the VPS via SSH or reads from
a local CSV for offline use.

Usage:
    python3 scripts/generate-report.py                    # Live VPS data
    python3 scripts/generate-report.py --dry-run          # Print summary only
    python3 scripts/generate-report.py --csv data.csv     # From local CSV
    python3 scripts/generate-report.py --date "March 15"  # Custom date label
    python3 scripts/generate-report.py --output report.docx

Dependencies: pandas, matplotlib, python-docx (pip install python-docx)
"""

import argparse
import io
import json
import os
import subprocess
import sys
from datetime import datetime, timezone

import matplotlib
matplotlib.use("Agg")  # Non-interactive backend
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips
from report_utils import format_int

# ---------------------------------------------------------------------------
# Constants & Design Tokens
# ---------------------------------------------------------------------------
# Aligned with web/src/styles/tokens.css, adapted for light-mode print.

ACCENT_COLOR = "#1083FE"
ACCENT_RGB = RGBColor(0x10, 0x83, 0xFE)
ACCENT_DARK_RGB = RGBColor(0x0A, 0x74, 0xE8)
LIGHT_ACCENT = "#EAF4FF"
CARD_BG = "#F8FAFC"
ROW_ALT_COLOR = "#F8FBFF"
BORDER_COLOR = "D7E6FB"
BORDER_COLOR_HEX = "#D7E6FB"
TEXT_PRIMARY_RGB = RGBColor(0x1F, 0x29, 0x37)
TEXT_SECONDARY_RGB = RGBColor(0x5F, 0x6B, 0x7A)

# Status colors (from tokens.css)
SUCCESS_HEX = "#34C759"
WARNING_HEX = "#FF9F0A"
ERROR_HEX = "#FF453A"

# Typography
FONT = "Segoe UI"
TITLE_SIZE = Pt(20)
H1_SIZE = Pt(14)
H2_SIZE = Pt(12)
BODY_SIZE = Pt(10)
SMALL_SIZE = Pt(8.5)
METHODOLOGY_SIZE = Pt(9)

# Chart constants
CHART_GRID_COLOR = "#E8EEF8"
CHART_TEXT_COLOR = "#5F6B7A"
CHART_SPINE_COLOR = "#D7E6FB"
CHART_DPI = 150

SCORE_COMPONENTS = [
    "recency_score", "engagement_score", "bridging_score",
    "source_diversity_score", "relevance_score",
]
COMPONENT_LABELS = ["Recency", "Engagement", "Bridging", "Source Diversity", "Relevance"]
WEIGHT_COLUMNS = [
    "recency_weight", "engagement_weight", "bridging_weight",
    "source_diversity_weight", "relevance_weight",
]

# SQL for the 3 data pulls (separated by marker in SSH output)
MARKER = "---REPORT-MARKER---"

POSTS_SQL = """
COPY (
  SELECT
    p.uri, LEFT(p.text, 200) as text, p.author_did,
    p.embed_url,
    p.topic_vector::text as topics,
    p.classification_method,
    ps.total_score,
    ps.recency_score, ps.engagement_score, ps.bridging_score,
    ps.source_diversity_score, ps.relevance_score,
    ps.recency_weight, ps.engagement_weight, ps.bridging_weight,
    ps.source_diversity_weight, ps.relevance_weight,
    ps.recency_weighted, ps.engagement_weighted, ps.bridging_weighted,
    ps.source_diversity_weighted, ps.relevance_weighted,
    COALESCE(pe.like_count,0) as likes,
    COALESCE(pe.repost_count,0) as reposts,
    COALESCE(pe.reply_count,0) as replies
  FROM post_scores ps
  JOIN posts p ON p.uri = ps.post_uri
  LEFT JOIN post_engagement pe ON pe.post_uri = ps.post_uri
  WHERE ps.epoch_id = (
    SELECT id FROM governance_epochs
    WHERE status='active' ORDER BY id DESC LIMIT 1
  )
    AND p.deleted = FALSE
  ORDER BY ps.total_score DESC
  LIMIT 1000
) TO STDOUT WITH CSV HEADER
""".strip().replace("\n", " ")

EPOCH_SQL = """
SELECT row_to_json(e) FROM (
  SELECT id, status,
    recency_weight, engagement_weight, bridging_weight,
    source_diversity_weight, relevance_weight,
    vote_count, topic_weights::text as topic_weights, created_at::text
  FROM governance_epochs
  WHERE status='active' ORDER BY id DESC LIMIT 1
) e
""".strip().replace("\n", " ")

STATS_SQL = """
SELECT row_to_json(s) FROM (
  SELECT
    (SELECT COUNT(*) FROM posts WHERE deleted=FALSE) as total_posts,
    (SELECT COUNT(*) FROM posts WHERE deleted=FALSE
     AND indexed_at > NOW() - INTERVAL '24 hours') as last_24h,
    (SELECT COUNT(*) FROM post_scores WHERE epoch_id = (
      SELECT id FROM governance_epochs
      WHERE status='active' ORDER BY id DESC LIMIT 1
    )) as scored_count
) s
""".strip().replace("\n", " ")


# ---------------------------------------------------------------------------
# Data Extraction
# ---------------------------------------------------------------------------

def fetch_from_vps():
    """Pull all three datasets from VPS in a single SSH call."""
    combined_cmd = (
        f'docker exec bluesky-feed-postgres psql -U feed -d bluesky_feed -c "{POSTS_SQL}" '
        f'&& echo "{MARKER}" '
        f'&& docker exec bluesky-feed-postgres psql -U feed -d bluesky_feed -t -A -c "{EPOCH_SQL}" '
        f'&& echo "{MARKER}" '
        f'&& docker exec bluesky-feed-postgres psql -U feed -d bluesky_feed -t -A -c "{STATS_SQL}"'
    )
    print("Connecting to VPS and pulling data...")
    result = subprocess.run(
        ["ssh", "corgi-vps", combined_cmd],
        capture_output=True, text=True, timeout=60,
    )
    if result.returncode != 0:
        print(f"SSH failed (exit {result.returncode}):", file=sys.stderr)
        print(result.stderr, file=sys.stderr)
        sys.exit(1)

    parts = result.stdout.split(MARKER)
    if len(parts) < 3:
        print(f"Expected 3 data sections, got {len(parts)}", file=sys.stderr)
        print("Raw output preview:", result.stdout[:500], file=sys.stderr)
        sys.exit(1)

    posts_csv = parts[0].strip()
    epoch_json_str = parts[1].strip()
    stats_json_str = parts[2].strip()

    df = pd.read_csv(io.StringIO(posts_csv))
    epoch = json.loads(epoch_json_str) if epoch_json_str else {}
    stats = json.loads(stats_json_str) if stats_json_str else {}

    # Parse topic_weights from nested string
    if "topic_weights" in epoch and isinstance(epoch["topic_weights"], str):
        try:
            epoch["topic_weights"] = json.loads(epoch["topic_weights"])
        except (json.JSONDecodeError, TypeError):
            epoch["topic_weights"] = {}

    print(f"  Posts: {len(df)}, Epoch: {epoch.get('id', '?')}, "
          f"Total indexed: {stats.get('total_posts', '?')}")
    return df, epoch, stats


def load_from_csv(csv_path, epoch_json_path=None):
    """Load data from local CSV file."""
    df = pd.read_csv(csv_path)
    epoch = {}
    stats = {"total_posts": len(df), "last_24h": 0, "scored_count": len(df)}
    if epoch_json_path and os.path.exists(epoch_json_path):
        with open(epoch_json_path, encoding="utf-8") as f:
            epoch = json.load(f)
    return df, epoch, stats


def parse_topic_vector(tv_str):
    """Parse a topic_vector JSONB string into a dict."""
    if not tv_str or pd.isna(tv_str) or tv_str in ("", "null", "None"):
        return {}
    try:
        return json.loads(tv_str.replace("'", '"'))
    except (json.JSONDecodeError, TypeError):
        return {}


def extract_primary_topic(tv_str):
    """Get the highest-confidence topic from a topic_vector."""
    tv = parse_topic_vector(tv_str)
    if not tv:
        return "uncategorized"
    return max(tv, key=tv.get)


def extract_topic_confidence(tv_str):
    """Get the confidence of the primary topic."""
    tv = parse_topic_vector(tv_str)
    if not tv:
        return 0.0
    return max(tv.values())


# ---------------------------------------------------------------------------
# Matplotlib Setup & Chart Generation
# ---------------------------------------------------------------------------

def setup_matplotlib_defaults():
    """Configure matplotlib for consistent, professional chart output."""
    plt.rcParams.update({
        "font.family": "sans-serif",
        "font.sans-serif": ["Segoe UI", "Helvetica Neue", "Helvetica", "Arial", "sans-serif"],
        "font.size": 9,
        "axes.titlesize": 11,
        "axes.titleweight": "semibold",
        "axes.labelsize": 9,
        "axes.labelcolor": CHART_TEXT_COLOR,
        "xtick.labelsize": 8,
        "ytick.labelsize": 8,
        "xtick.color": CHART_TEXT_COLOR,
        "ytick.color": CHART_TEXT_COLOR,
        "figure.facecolor": "white",
        "axes.facecolor": "white",
        "savefig.dpi": CHART_DPI,
        "savefig.bbox": "tight",
    })


def configure_chart_axes(ax):
    """Apply consistent axis styling to a chart axes object."""
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["bottom"].set_color(CHART_SPINE_COLOR)
    ax.spines["left"].set_color(CHART_SPINE_COLOR)
    ax.tick_params(colors=CHART_TEXT_COLOR, labelsize=8)
    ax.set_axisbelow(True)


def _save_chart(fig):
    """Save a matplotlib figure to PNG bytes and clean up."""
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=CHART_DPI, bbox_inches="tight",
                facecolor="white", edgecolor="none")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def generate_topic_chart(df):
    """Horizontal bar chart of top 15 topics with value labels."""
    topics = df["primary_topic"].value_counts().head(15)
    sorted_topics = topics.sort_values()

    fig, ax = plt.subplots(figsize=(7.5, 4.5))
    bars = ax.barh(range(len(sorted_topics)), sorted_topics.values,
                   color=ACCENT_COLOR, edgecolor="white", height=0.7)
    ax.set_yticks(range(len(sorted_topics)))
    ax.set_yticklabels([t.replace("-", " ").title() for t in sorted_topics.index],
                       fontsize=8)
    ax.set_xlabel("Post Count", fontsize=9, color=CHART_TEXT_COLOR)
    ax.set_title("Top 15 Topics in Feed", fontsize=11, fontweight="semibold",
                 color="#1F2937", pad=12)
    configure_chart_axes(ax)
    ax.grid(axis="x", color=CHART_GRID_COLOR, linewidth=0.6, alpha=0.8)
    ax.bar_label(bars, padding=4, fontsize=7, color=CHART_TEXT_COLOR, fmt="%d")

    # Add a bit of right margin for labels
    x_max = sorted_topics.max()
    ax.set_xlim(right=x_max * 1.12)
    fig.tight_layout()
    return _save_chart(fig)


def generate_score_boxplot(df):
    """Box plot of the 5 score components with median annotations."""
    data = [df[col].dropna() for col in SCORE_COMPONENTS]
    fig, ax = plt.subplots(figsize=(8.5, 4.0))

    bp = ax.boxplot(
        data, tick_labels=COMPONENT_LABELS, patch_artist=True,
        medianprops=dict(color="white", linewidth=1.5),
        flierprops=dict(marker=".", markersize=3, markerfacecolor=CHART_TEXT_COLOR,
                        markeredgecolor=CHART_TEXT_COLOR, alpha=0.4),
        whiskerprops=dict(color=CHART_TEXT_COLOR, linewidth=0.8),
        capprops=dict(color=CHART_TEXT_COLOR, linewidth=0.8),
    )
    for patch in bp["boxes"]:
        patch.set_facecolor(ACCENT_COLOR)
        patch.set_alpha(0.85)
        patch.set_edgecolor(ACCENT_COLOR)

    # Median value annotations
    for i, line in enumerate(bp["medians"]):
        x, y = line.get_xydata()[1]
        median_val = data[i].median()
        ax.text(x + 0.15, y, f"{median_val:.3f}",
                fontsize=7, color=CHART_TEXT_COLOR, va="center", fontweight="medium")

    ax.set_ylabel("Raw Score (0\u20131)", fontsize=9, color=CHART_TEXT_COLOR)
    ax.set_title("Score Component Distribution", fontsize=11,
                 fontweight="semibold", color="#1F2937", pad=12)
    configure_chart_axes(ax)
    ax.grid(axis="y", color=CHART_GRID_COLOR, linewidth=0.6, alpha=0.8)
    fig.tight_layout()
    return _save_chart(fig)


def generate_classification_donut(df):
    """Donut chart with colorblind-safe palette and center total."""
    counts = df["classification_method"].fillna("unknown").value_counts()
    labels = [m.capitalize() for m in counts.index]
    # Colorblind-safe: blue, green, amber (from design tokens status colors)
    palette = [ACCENT_COLOR, SUCCESS_HEX, WARNING_HEX]
    colors = palette[:len(counts)]

    fig, ax = plt.subplots(figsize=(4.2, 4.2))
    wedges, texts, autotexts = ax.pie(
        counts.values, labels=labels, autopct="%1.1f%%",
        colors=colors, startangle=90, pctdistance=0.75,
        wedgeprops=dict(width=0.4, edgecolor="white", linewidth=2),
    )
    for t in autotexts:
        t.set_fontsize(9)
        t.set_fontweight("bold")
        t.set_color("#1F2937")
    for t in texts:
        t.set_fontsize(9)
        t.set_color(CHART_TEXT_COLOR)

    # Center label — total count
    total = int(counts.sum())
    ax.text(0, 0.05, f"{total:,}", ha="center", va="center",
            fontsize=16, fontweight="bold", color="#1F2937")
    ax.text(0, -0.12, "posts", ha="center", va="center",
            fontsize=8, color=CHART_TEXT_COLOR)

    ax.set_title("Classification Method", fontsize=11,
                 fontweight="semibold", color="#1F2937", pad=12)
    fig.tight_layout()
    return _save_chart(fig)


# ---------------------------------------------------------------------------
# Document Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, hex_color):
    """Apply background shading to a table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color.replace("#", ""))
    shading.set(qn("w:val"), "clear")
    tc_pr.append(shading)


def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    """Set cell margins in twips via XML."""
    tc_pr = cell._element.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcMar"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_mar = OxmlElement("w:tcMar")
    for edge, val in [("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tc_mar.append(el)
    tc_pr.append(tc_mar)


def set_table_borders(table, color=BORDER_COLOR, size=4, inside_v=True):
    """Apply subtle table borders for a cleaner, print-friendly look."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    edges = ["top", "left", "bottom", "right", "insideH"]
    if inside_v:
        edges.append("insideV")

    for edge in edges:
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_row_border(row, edge="bottom", color=BORDER_COLOR, size=6):
    """Set a specific border on all cells in a row."""
    for cell in row.cells:
        tc_pr = cell._element.get_or_add_tcPr()
        tc_borders = tc_pr.find(qn("w:tcBorders"))
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)
        el = tc_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tc_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def add_paragraph_border(paragraph, edge="bottom", color=BORDER_COLOR, size=4, space=4):
    """Add a border to a paragraph via XML."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), color)
    p_bdr.append(el)


def add_simple_field(paragraph, field_code):
    """Insert a Word field (e.g., PAGE or NUMPAGES) into a paragraph."""
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), field_code)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    fld.append(run)
    paragraph._p.append(fld)


def configure_document_styles(doc):
    """Configure reusable typography styles for a polished report."""
    # Normal body style
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = BODY_SIZE
    normal.font.color.rgb = TEXT_PRIMARY_RGB
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.widow_control = True

    # Title
    title_style = doc.styles["Title"]
    title_style.font.name = FONT
    title_style.font.size = TITLE_SIZE
    title_style.font.bold = True
    title_style.font.color.rgb = ACCENT_DARK_RGB
    title_style.paragraph_format.space_before = Pt(48)
    title_style.paragraph_format.space_after = Pt(8)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Heading 1 — with bottom accent border
    h1 = doc.styles["Heading 1"]
    h1.font.name = FONT
    h1.font.size = H1_SIZE
    h1.font.bold = True
    h1.font.color.rgb = ACCENT_DARK_RGB
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = FONT
    h2.font.size = H2_SIZE
    h2.font.bold = True
    h2.font.color.rgb = ACCENT_RGB
    h2.paragraph_format.space_before = Pt(16)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    # Custom styles
    existing = {s.name for s in doc.styles}

    if "ReportCaption" not in existing:
        cap = doc.styles.add_style("ReportCaption", WD_STYLE_TYPE.PARAGRAPH)
        cap.font.name = FONT
        cap.font.size = SMALL_SIZE
        cap.font.color.rgb = TEXT_SECONDARY_RGB
        cap.font.italic = True
        cap.paragraph_format.space_before = Pt(4)
        cap.paragraph_format.space_after = Pt(10)
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if "ReportSubtitle" not in existing:
        sub = doc.styles.add_style("ReportSubtitle", WD_STYLE_TYPE.PARAGRAPH)
        sub.font.name = FONT
        sub.font.size = Pt(12)
        sub.font.color.rgb = TEXT_SECONDARY_RGB
        sub.paragraph_format.space_after = Pt(16)
        sub.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if "MethodologyBody" not in existing:
        meth = doc.styles.add_style("MethodologyBody", WD_STYLE_TYPE.PARAGRAPH)
        meth.font.name = FONT
        meth.font.size = METHODOLOGY_SIZE
        meth.font.color.rgb = TEXT_SECONDARY_RGB
        meth.paragraph_format.space_after = Pt(4)
        meth.paragraph_format.line_spacing = 1.3

    if "AppendixHeading" not in existing:
        app_h = doc.styles.add_style("AppendixHeading", WD_STYLE_TYPE.PARAGRAPH)
        app_h.font.name = FONT
        app_h.font.size = Pt(13)
        app_h.font.bold = True
        app_h.font.color.rgb = ACCENT_RGB
        app_h.paragraph_format.space_before = Pt(20)
        app_h.paragraph_format.space_after = Pt(8)
        app_h.paragraph_format.keep_with_next = True


def configure_section_layout(section, landscape=False):
    """Apply consistent page layout and orientation."""
    if landscape and section.orientation != WD_ORIENT.LANDSCAPE:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    if not landscape and section.orientation != WD_ORIENT.PORTRAIT:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width, section.page_height = section.page_height, section.page_width

    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)


def add_landscape_section(doc, date_label, epoch_id):
    """Insert a new landscape section with header/footer."""
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section_layout(section, landscape=True)
    set_document_header_footer(section, date_label, epoch_id)
    return section


def restore_portrait_section(doc, date_label, epoch_id):
    """Insert a new portrait section with header/footer."""
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section_layout(section, landscape=False)
    set_document_header_footer(section, date_label, epoch_id)
    return section


def set_document_header_footer(section, date_label, epoch_id):
    """Set professional header/footer with borders and page numbering."""
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    # --- Header: left title, right date, bottom border ---
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = ""
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Left-aligned title
    title_run = header_para.add_run("Corgi Network Feed Quality Report")
    title_run.font.name = FONT
    title_run.font.size = SMALL_SIZE
    title_run.font.color.rgb = TEXT_SECONDARY_RGB

    # Tab + right-aligned date
    tab_run = header_para.add_run("\t\t")
    tab_run.font.name = FONT
    tab_run.font.size = SMALL_SIZE
    date_run = header_para.add_run(date_label)
    date_run.font.name = FONT
    date_run.font.size = SMALL_SIZE
    date_run.font.color.rgb = TEXT_SECONDARY_RGB

    # Add right tab stop
    is_landscape = section.orientation == WD_ORIENT.LANDSCAPE
    tab_pos = Inches(9.5) if is_landscape else Inches(7.0)
    pPr = header_para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(int(tab_pos)))
    tabs.append(tab)
    pPr.append(tabs)

    # Bottom border on header
    add_paragraph_border(header_para, "bottom", BORDER_COLOR, size=4, space=6)

    # --- Footer: feed.corgi.network (left), Epoch | Page (center), top border ---
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = ""
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Top border on footer
    add_paragraph_border(footer_para, "top", BORDER_COLOR, size=4, space=6)

    # Footer content
    site_run = footer_para.add_run("feed.corgi.network  |  ")
    site_run.font.name = FONT
    site_run.font.size = Pt(7.5)
    site_run.font.color.rgb = TEXT_SECONDARY_RGB

    meta = footer_para.add_run(f"Epoch {epoch_id}  |  ")
    meta.font.name = FONT
    meta.font.size = SMALL_SIZE
    meta.font.color.rgb = TEXT_SECONDARY_RGB

    page_label = footer_para.add_run("Page ")
    page_label.font.name = FONT
    page_label.font.size = SMALL_SIZE
    page_label.font.color.rgb = TEXT_SECONDARY_RGB
    add_simple_field(footer_para, "PAGE")

    of_run = footer_para.add_run(" of ")
    of_run.font.name = FONT
    of_run.font.size = SMALL_SIZE
    of_run.font.color.rgb = TEXT_SECONDARY_RGB
    add_simple_field(footer_para, "NUMPAGES")


def styled_heading(doc, text, level=1, style_name=None):
    """Add a styled heading. Use style_name='AppendixHeading' for appendix sections."""
    if style_name:
        heading = doc.add_paragraph(text, style=style_name)
    else:
        heading = doc.add_heading(text, level=level)
    # Add bottom accent border to Heading 1
    if level == 1 and not style_name:
        add_paragraph_border(heading, "bottom", BORDER_COLOR, size=4, space=4)
    return heading


def styled_paragraph(doc, text, bold=False, size=None, style=None):
    """Add a styled body paragraph."""
    p = doc.add_paragraph(style=style or "Normal")
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = size or BODY_SIZE
    run.bold = bold
    run.font.color.rgb = TEXT_PRIMARY_RGB
    return p


def styled_caption(doc, text):
    """Add a centered caption under charts/figures."""
    return doc.add_paragraph(text, style="ReportCaption")


def add_callout_box(doc, bullets, label="Key Takeaways"):
    """Add a styled callout box with multi-bullet support and left accent border."""
    # Backward compat: string → list
    if isinstance(bullets, str):
        bullets = [bullets]

    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Left accent border only — clean "quote block" look
    tbl = box._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ["top", "right", "bottom", "insideH", "insideV"]:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    left_border = OxmlElement("w:left")
    left_border.set(qn("w:val"), "single")
    left_border.set(qn("w:sz"), "18")  # ~3pt
    left_border.set(qn("w:space"), "0")
    left_border.set(qn("w:color"), ACCENT_COLOR.replace("#", ""))
    borders.append(left_border)
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(borders)

    cell = box.rows[0].cells[0]
    set_cell_shading(cell, CARD_BG)
    set_cell_margins(cell, top=100, bottom=100, left=180, right=140)

    # Label paragraph
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    lead = p.add_run(label)
    lead.bold = True
    lead.font.name = FONT
    lead.font.size = BODY_SIZE
    lead.font.color.rgb = ACCENT_DARK_RGB

    # Bullet paragraphs
    for bullet_text in bullets:
        bp = cell.add_paragraph()
        bp.paragraph_format.space_before = Pt(2)
        bp.paragraph_format.space_after = Pt(2)
        bp.paragraph_format.line_spacing = 1.3
        run = bp.add_run(f"\u2022  {bullet_text}")
        run.font.name = FONT
        run.font.size = SMALL_SIZE
        run.font.color.rgb = TEXT_PRIMARY_RGB

    # Space after the callout via paragraph spacing on the wrapper
    wrapper = doc.paragraphs[-1] if doc.paragraphs else None
    if wrapper:
        wrapper.paragraph_format.space_after = Pt(12)


def add_styled_table(doc, headers, rows, col_widths=None, numeric_cols=None):
    """Add a formatted table with accent header, subtle borders, and cell padding."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    numeric_cols = numeric_cols or set()

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, ACCENT_COLOR)
        set_cell_margins(cell, top=50, bottom=50, left=60, right=60)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = FONT
                run.font.size = SMALL_SIZE
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True

    # Thicker bottom border on header row
    set_row_border(table.rows[0], "bottom",
                   color=ACCENT_COLOR.replace("#", ""), size=8)

    # Column widths
    if col_widths:
        for row in table.rows:
            for c, width in enumerate(col_widths):
                row.cells[c].width = width

    # Data rows
    for r, row_data in enumerate(rows):
        for c, value in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(value)
            set_cell_margins(cell, top=40, bottom=40, left=60, right=60)
            if r % 2 == 1:
                set_cell_shading(cell, ROW_ALT_COLOR)
            for paragraph in cell.paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.RIGHT if c in numeric_cols
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                for run in paragraph.runs:
                    run.font.name = FONT
                    run.font.size = SMALL_SIZE
                    run.font.color.rgb = TEXT_PRIMARY_RGB

    return table


def add_horizontal_rule(doc):
    """Add a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    add_paragraph_border(p, "bottom", BORDER_COLOR, size=4, space=0)
    return p


# ---------------------------------------------------------------------------
# Table of Contents
# ---------------------------------------------------------------------------

def add_table_of_contents(doc):
    """Insert a Word TOC field that can be updated when opened."""
    styled_heading(doc, "Contents")

    # TOC field
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)

    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr_run = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-2" \\h \\z \\u '
    instr_run._r.append(instr_text)

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    sep_run = paragraph.add_run()
    sep_run._r.append(fld_char_separate)

    # Placeholder text
    placeholder = paragraph.add_run(
        "[Right-click \u2192 Update Field to populate table of contents]"
    )
    placeholder.font.name = FONT
    placeholder.font.size = SMALL_SIZE
    placeholder.font.color.rgb = TEXT_SECONDARY_RGB
    placeholder.font.italic = True

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    end_run = paragraph.add_run()
    end_run._r.append(fld_char_end)


# ---------------------------------------------------------------------------
# Page Builders
# ---------------------------------------------------------------------------

def build_title_page(doc, date_label, epoch):
    """Title page with report name and subtitle."""
    epoch_id = epoch.get("id", "?")

    title = doc.add_paragraph(style="Title")
    run = title.add_run(f"Feed Quality Analysis")
    run.font.name = FONT

    # Date subtitle
    date_p = doc.add_paragraph(style="ReportSubtitle")
    run = date_p.add_run(date_label)
    run.font.name = FONT

    # Info line
    info = doc.add_paragraph(style="ReportSubtitle")
    run = info.add_run(
        f"Community-Governed Feed  |  feed.corgi.network  |  Epoch {epoch_id}"
    )
    run.font.name = FONT
    run.font.size = BODY_SIZE


def build_executive_summary(doc, df, epoch, stats):
    """Executive summary with key metrics and governance state."""
    styled_heading(doc, "Executive Summary")

    # 4-stat banner
    unique_authors = df["author_did"].nunique() if "author_did" in df.columns else 0
    median_score = df["total_score"].median() if "total_score" in df.columns else 0
    embedding_pct = 0
    if "classification_method" in df.columns:
        emb_count = (df["classification_method"] == "embedding").sum()
        embedding_pct = (emb_count / len(df) * 100) if len(df) > 0 else 0

    banner_data = [
        ("Posts in Scope", format_int(len(df))),
        ("Unique Authors", format_int(unique_authors)),
        ("Median Score", f"{median_score:.3f}"),
        ("% Embedding", f"{embedding_pct:.1f}%"),
    ]
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, size=4, inside_v=True)
    for i, (label, value) in enumerate(banner_data):
        # Value row (large number)
        cell = table.rows[0].cells[i]
        cell.text = value
        set_cell_shading(cell, ACCENT_COLOR)
        set_cell_margins(cell, top=80, bottom=60, left=40, right=40)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = FONT
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
        # Label row
        cell = table.rows[1].cells[i]
        cell.text = label
        set_cell_shading(cell, LIGHT_ACCENT)
        set_cell_margins(cell, top=40, bottom=40, left=40, right=40)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = FONT
                run.font.size = SMALL_SIZE
                run.font.color.rgb = ACCENT_RGB

    # Summary paragraph
    total_posts = stats.get("total_posts", "?")
    scored_count = stats.get("scored_count", len(df))
    topic_count = df["primary_topic"].nunique() if "primary_topic" in df.columns else 0
    summary = (
        f"The top {len(df):,} feed posts span {format_int(unique_authors)} unique "
        f"authors across {format_int(topic_count)} topics. The database contains "
        f"{format_int(total_posts)} indexed posts with {format_int(scored_count)} "
        f"scored in the current epoch."
    )
    p = styled_paragraph(doc, summary)
    p.paragraph_format.space_before = Pt(12)

    # Key Takeaways
    top_topic = (
        df["primary_topic"].value_counts().index[0]
        if "primary_topic" in df.columns and len(df) > 0
        else "unknown"
    )
    add_callout_box(doc, [
        f"Most represented topic: {top_topic.replace('-', ' ').title()}",
        f"{embedding_pct:.1f}% of analyzed posts used embedding classification",
        f"{format_int(unique_authors)} unique authors contribute to feed diversity",
    ])

    # Governance weights table
    styled_heading(doc, "Epoch Weights", level=2)
    weight_names = ["Recency", "Engagement", "Bridging", "Source Diversity", "Relevance"]
    weight_keys = ["recency_weight", "engagement_weight", "bridging_weight",
                   "source_diversity_weight", "relevance_weight"]
    weight_rows = []
    for name, key in zip(weight_names, weight_keys):
        val = epoch.get(key, "?")
        weight_rows.append((name, f"{val:.2f}" if isinstance(val, (int, float)) else str(val)))
    weight_rows.append(("Vote Count", format_int(epoch.get("vote_count", "?"))))
    add_styled_table(
        doc,
        ["Parameter", "Value"],
        weight_rows,
        col_widths=[Inches(4.7), Inches(2.0)],
        numeric_cols={1},
    )


def build_methodology_block(doc, epoch, stats, date_label):
    """Data freshness and methodology disclosure block."""
    styled_heading(doc, "Methodology & Data Freshness", level=2)

    # Methodology content in a subtle box
    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(box, color=BORDER_COLOR, size=4)
    cell = box.rows[0].cells[0]
    set_cell_shading(cell, CARD_BG)
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)

    epoch_id = epoch.get("id", "?")
    created_at = epoch.get("created_at", "unknown")
    total_posts = format_int(stats.get("total_posts", "?"))
    scored_count = format_int(stats.get("scored_count", "?"))
    last_24h = format_int(stats.get("last_24h", "?"))
    gen_time = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    lines = [
        f"Scope: Top 1,000 scored posts from the active governance epoch.",
        f"Epoch: {epoch_id} (created {created_at}).",
        f"Database: {total_posts} total indexed posts, {scored_count} scored this epoch, "
        f"{last_24h} ingested in last 24 hours.",
        f"Generated: {gen_time}.",
        f"Note: Engagement metrics (likes, reposts, replies) reflect counts at time of "
        f"data extraction and may differ from current values.",
    ]

    p = cell.paragraphs[0]
    p.text = ""
    for i, line in enumerate(lines):
        if i > 0:
            p = cell.add_paragraph()
        p.style = doc.styles["MethodologyBody"]
        run = p.add_run(line)
        run.font.name = FONT
        run.font.size = METHODOLOGY_SIZE
        run.font.color.rgb = TEXT_SECONDARY_RGB


def build_topic_weights_page(doc, epoch):
    """Dedicated topic-weight table to avoid overflow on executive page."""
    topic_weights = epoch.get("topic_weights", {})
    if not (topic_weights and isinstance(topic_weights, dict)):
        return False

    styled_heading(doc, "Topic Weights")
    styled_paragraph(
        doc,
        "Topic-level governance weights for the currently active epoch.",
    )
    sorted_tw = sorted(topic_weights.items(), key=lambda x: x[1], reverse=True)[:10]
    tw_rows = [(slug.replace("-", " ").title(), f"{w:.3f}") for slug, w in sorted_tw]
    add_styled_table(
        doc,
        ["Topic", "Weight"],
        tw_rows,
        col_widths=[Inches(4.7), Inches(2.0)],
        numeric_cols={1},
    )
    return True


def build_topic_page(doc, df):
    """Topic distribution chart and stats table."""
    styled_heading(doc, "Topic Distribution")

    # Chart
    chart_bytes = generate_topic_chart(df)
    doc.add_picture(io.BytesIO(chart_bytes), width=Inches(6.6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_caption(doc, "Figure 1. Top topics among the top 1,000 scored posts.")

    # Stats table
    topic_stats = df.groupby("primary_topic").agg(
        count=("total_score", "size"),
        avg_relevance=("relevance_score", "mean"),
    ).sort_values("count", ascending=False).head(15)

    total = len(df)
    rows = []
    for topic, row in topic_stats.iterrows():
        pct = (row["count"] / total * 100) if total > 0 else 0
        rows.append((
            str(topic).replace("-", " ").title(),
            format_int(row["count"]),
            f"{pct:.1f}%",
            f"{row['avg_relevance']:.3f}",
        ))
    add_styled_table(
        doc,
        ["Topic", "Post Count", "% of Feed", "Avg Relevance"],
        rows,
        col_widths=[Inches(3.3), Inches(1.1), Inches(1.1), Inches(1.2)],
        numeric_cols={1, 2, 3},
    )

    # Key Takeaways
    top3_pct = (
        df["primary_topic"].value_counts().head(3).sum() / len(df) * 100
        if len(df) > 0 and "primary_topic" in df.columns
        else 0
    )
    top_relevance_topic = (
        topic_stats["avg_relevance"].idxmax()
        if len(topic_stats) > 0 else "unknown"
    )
    top_relevance_val = (
        topic_stats["avg_relevance"].max()
        if len(topic_stats) > 0 else 0
    )
    total_topics = df["primary_topic"].nunique() if "primary_topic" in df.columns else 0

    add_callout_box(doc, [
        f"Top 3 topics account for {top3_pct:.1f}% of the analyzed feed sample",
        f"Highest avg relevance: {str(top_relevance_topic).replace('-', ' ').title()} "
        f"({top_relevance_val:.3f})",
        f"{format_int(total_topics)} distinct topics detected across the sample",
    ])


def build_score_page(doc, df):
    """Score composition box plot and stats (rendered on landscape page)."""
    styled_heading(doc, "Score Composition")

    # Box plot chart
    chart_bytes = generate_score_boxplot(df)
    doc.add_picture(io.BytesIO(chart_bytes), width=Inches(8.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_caption(doc, "Figure 2. Distribution of raw score components before weighting.")

    # Stats table — wider columns for landscape
    headers = ["Component", "Min", "P25", "Median", "P75", "Max", "Weight"]
    rows = []
    for col, label, wcol in zip(SCORE_COMPONENTS, COMPONENT_LABELS, WEIGHT_COLUMNS):
        series = df[col].dropna()
        weight_val = df[wcol].iloc[0] if wcol in df.columns and len(df) > 0 else "?"
        rows.append((
            label,
            f"{series.min():.3f}" if len(series) > 0 else "N/A",
            f"{series.quantile(0.25):.3f}" if len(series) > 0 else "N/A",
            f"{series.median():.3f}" if len(series) > 0 else "N/A",
            f"{series.quantile(0.75):.3f}" if len(series) > 0 else "N/A",
            f"{series.max():.3f}" if len(series) > 0 else "N/A",
            f"{weight_val:.2f}" if isinstance(weight_val, (int, float)) else str(weight_val),
        ))
    add_styled_table(
        doc,
        headers,
        rows,
        col_widths=[
            Inches(2.0), Inches(1.2), Inches(1.2), Inches(1.2),
            Inches(1.2), Inches(1.2), Inches(1.1),
        ],
        numeric_cols={1, 2, 3, 4, 5, 6},
    )

    # Key Takeaways
    medians = {label: df[col].median() for col, label in zip(SCORE_COMPONENTS, COMPONENT_LABELS)}
    highest = max(medians, key=medians.get)
    lowest = min(medians, key=medians.get)
    gap = medians[highest] - medians[lowest]

    # Find which weight is highest
    weight_vals = {}
    for label, wcol in zip(COMPONENT_LABELS, WEIGHT_COLUMNS):
        if wcol in df.columns and len(df) > 0:
            weight_vals[label] = df[wcol].iloc[0]
    top_weight = max(weight_vals, key=weight_vals.get) if weight_vals else "?"
    top_weight_val = weight_vals.get(top_weight, 0)

    add_callout_box(doc, [
        f"{highest} has the highest median raw score ({medians[highest]:.3f}), "
        f"{lowest} the lowest ({medians[lowest]:.3f})",
        f"Median gap between highest and lowest components: {gap:.3f}",
        f"Highest governance weight: {top_weight} ({top_weight_val:.2f})",
    ])


def build_classification_page(doc, df):
    """Classification method distribution."""
    styled_heading(doc, "Classification Methods")

    # Donut chart
    chart_bytes = generate_classification_donut(df)
    doc.add_picture(io.BytesIO(chart_bytes), width=Inches(4.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_caption(doc, "Figure 3. Classification method share in the analyzed sample.")

    # Classification stats
    method_counts = df["classification_method"].fillna("unknown").value_counts()
    method_rows = [(m.capitalize(), format_int(c), f"{c/len(df)*100:.1f}%")
                   for m, c in method_counts.items()]
    add_styled_table(
        doc,
        ["Method", "Count", "% of Feed"],
        method_rows,
        col_widths=[Inches(3.3), Inches(1.2), Inches(1.4)],
        numeric_cols={1, 2},
    )

    # Key Takeaways
    dominant_method = method_counts.index[0] if len(method_counts) > 0 else "unknown"
    dominant_pct = (method_counts.iloc[0] / len(df) * 100) if len(df) > 0 else 0
    method_count = len(method_counts)
    add_callout_box(doc, [
        f"{str(dominant_method).capitalize()} is the dominant classifier at "
        f"{dominant_pct:.1f}% of analyzed posts",
        f"{method_count} classification method(s) active in the current pipeline",
    ])


def build_diversity_page(doc, df):
    """Author diversity and URL deduplication diagnostics."""
    styled_heading(doc, "Author Diversity & URL Deduplication")

    # Author diversity
    author_counts = df["author_did"].dropna().value_counts()
    total_unique_authors = df["author_did"].dropna().nunique()
    avg_posts_per_author = author_counts.mean() if len(author_counts) > 0 else 0
    max_posts_single_author = int(author_counts.max()) if len(author_counts) > 0 else 0
    styled_paragraph(
        doc,
        f"Total unique authors: {format_int(total_unique_authors)}. "
        f"Average posts per author: {avg_posts_per_author:.1f}. "
        f"Max posts by single author: {format_int(max_posts_single_author)}.",
    )

    # Flag authors with 5+ posts
    heavy_posters = author_counts[author_counts >= 5]
    if len(heavy_posters) > 0:
        rows = [(did[:20] + "...", format_int(count))
                for did, count in heavy_posters.head(10).items()]
        add_styled_table(
            doc,
            ["Author DID (truncated)", "Post Count"],
            rows,
            col_widths=[Inches(4.7), Inches(2.0)],
            numeric_cols={1},
        )

    # URL dedup stats
    max_duplicates = 0
    if "embed_url" in df.columns:
        styled_heading(doc, "URL Deduplication", level=2)
        normalized_urls = (
            df["embed_url"]
            .dropna()
            .astype(str)
            .str.strip()
        )
        normalized_urls = normalized_urls[~normalized_urls.isin(["", "null", "None"])]

        url_counts = normalized_urls.value_counts()
        duplicate_urls = url_counts[url_counts > 1]
        duplicate_posts = int(duplicate_urls.sum()) if len(duplicate_urls) > 0 else 0
        max_duplicates = int(duplicate_urls.max()) if len(duplicate_urls) > 0 else 0

        styled_paragraph(
            doc,
            f"Posts with an embed URL: {format_int(len(normalized_urls))}. "
            f"Posts sharing a duplicated embed URL: {format_int(duplicate_posts)} "
            f"across {format_int(len(duplicate_urls))} URLs. "
            f"Max duplicates for a single URL: {format_int(max_duplicates)}.",
        )

    # Key Takeaways
    max_author_pct = (max_posts_single_author / len(df) * 100) if len(df) > 0 else 0
    takeaways = [
        f"Top author concentration: {max_posts_single_author} posts "
        f"({max_author_pct:.1f}% of sample)",
        f"{len(heavy_posters)} author(s) have 5+ posts in the top 1,000",
    ]
    if max_duplicates > 0:
        takeaways.append(
            f"Most repeated external URL appears {format_int(max_duplicates)} times"
        )
    add_callout_box(doc, takeaways)


def build_engagement_page(doc, df):
    """Engagement profile - top vs bottom comparison."""
    styled_heading(doc, "Engagement Profile")

    if len(df) < 200:
        styled_paragraph(doc, f"Insufficient data for top/bottom comparison "
                         f"(need 200+ posts, have {len(df)}).")
        return

    top100 = df.head(100)
    bottom100 = df.tail(100)

    metrics = ["likes", "reposts", "replies"]
    rows = []
    ratios = {}
    for metric in metrics:
        if metric not in df.columns:
            continue
        top_med = top100[metric].median()
        bot_med = bottom100[metric].median()
        ratio_val = top_med / bot_med if bot_med > 0 else float("inf")
        ratio_str = f"{ratio_val:.1f}x" if bot_med > 0 else "\u221e"
        ratios[metric] = ratio_str
        rows.append((
            metric.capitalize(),
            f"{top_med:.1f}",
            f"{bot_med:.1f}",
            ratio_str,
        ))

    add_styled_table(
        doc,
        ["Metric", "Top 100 Median", "Bottom 100 Median", "Ratio"],
        rows,
        col_widths=[Inches(2.0), Inches(1.5), Inches(1.7), Inches(1.3)],
        numeric_cols={1, 2, 3},
    )

    # Key Takeaways
    if "likes" in df.columns:
        top_likes = top100["likes"].median()
        bot_likes = bottom100["likes"].median()
        ratio = (top_likes / bot_likes) if bot_likes > 0 else 0

        takeaways = [
            f"Top-ranked posts receive ~{ratio:.1f}x median likes vs bottom-ranked",
        ]
        if "reposts" in ratios:
            takeaways.append(
                f"Repost ratio: {ratios['reposts']} (top 100 vs bottom 100)"
            )
        takeaways.append(
            "Scoring algorithm effectively surfaces community-resonant content"
        )
        add_callout_box(doc, takeaways)


def build_sample_posts_page(doc, subset, title, intro_text, is_appendix=False):
    """Sample post table, optionally styled as appendix (landscape expected)."""
    if is_appendix:
        add_horizontal_rule(doc)
        styled_heading(doc, title, style_name="AppendixHeading")
        styled_paragraph(doc, intro_text, size=METHODOLOGY_SIZE)
    else:
        styled_heading(doc, title)
        styled_paragraph(doc, intro_text)

    rows = []
    for _, row in subset.iterrows():
        text = str(row.get("text", ""))[:100]
        if len(str(row.get("text", ""))) > 100:
            text += "..."
        topic = str(row.get("primary_topic", "?")).replace("-", " ").title()
        score = f"{row.get('total_score', 0):.3f}"
        likes = format_int(int(row.get("likes", 0)))
        method = str(row.get("classification_method", "?")).capitalize()
        rows.append((text, topic, score, likes, method))

    # Landscape widths
    add_styled_table(
        doc,
        ["Text Preview", "Topic", "Score", "Likes", "Method"],
        rows,
        col_widths=[Inches(4.5), Inches(1.5), Inches(1.0), Inches(0.8), Inches(1.2)],
        numeric_cols={2, 3},
    )


# ---------------------------------------------------------------------------
# Dry Run
# ---------------------------------------------------------------------------

def print_dry_run(df, epoch, stats):
    """Print data summary to stdout without generating docx."""
    print("\n" + "=" * 60)
    print("DRY RUN - Data Summary")
    print("=" * 60)

    print(f"\nPosts loaded: {len(df)}")
    print(f"Unique authors: {df['author_did'].nunique()}")
    print(f"Median total score: {df['total_score'].median():.4f}")

    if "classification_method" in df.columns:
        print(f"\nClassification methods:")
        for method, count in df["classification_method"].fillna("unknown").value_counts().items():
            print(f"  {method}: {count} ({count/len(df)*100:.1f}%)")

    if "primary_topic" in df.columns:
        print(f"\nTop 10 topics:")
        for topic, count in df["primary_topic"].value_counts().head(10).items():
            print(f"  {topic}: {count}")

    print(f"\nScore components (medians):")
    for col, label in zip(SCORE_COMPONENTS, COMPONENT_LABELS):
        if col in df.columns:
            print(f"  {label}: {df[col].median():.4f}")

    if "likes" in df.columns:
        print(f"\nEngagement (medians): likes={df['likes'].median():.0f}, "
              f"reposts={df['reposts'].median():.0f}, "
              f"replies={df['replies'].median():.0f}")

    print(f"\nEpoch: {epoch.get('id', '?')} (status: {epoch.get('status', '?')}, "
          f"votes: {epoch.get('vote_count', '?')})")
    print(f"System stats: {stats.get('total_posts', '?')} total posts, "
          f"{stats.get('last_24h', '?')} last 24h, "
          f"{stats.get('scored_count', '?')} scored")
    print("\n" + "=" * 60)
    print("Dry run complete. No docx generated.")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Generate feed quality analysis report (.docx)",
    )
    parser.add_argument("--csv", help="Read from local CSV instead of VPS")
    parser.add_argument("--epoch-json", help="Epoch data as JSON file (for offline mode)")
    parser.add_argument("--output", help="Output file path")
    parser.add_argument("--date", help="Date label for report title")
    parser.add_argument("--dry-run", action="store_true",
                        help="Print data summary without generating docx")
    args = parser.parse_args()

    # Date label
    date_label = args.date or datetime.now().strftime("%B %d, %Y")

    # Default output path
    date_slug = datetime.now().strftime("%b%d").lower()
    default_output = f"reports/sprint-data-analysis-{date_slug}.docx"
    output_path = args.output or default_output

    # Data extraction
    if args.csv:
        print(f"Loading from CSV: {args.csv}")
        df, epoch, stats = load_from_csv(args.csv, args.epoch_json)
    else:
        df, epoch, stats = fetch_from_vps()

    if len(df) == 0:
        print("No data returned. Check VPS connectivity and epoch state.",
              file=sys.stderr)
        sys.exit(1)

    # Derived columns
    df["primary_topic"] = df["topics"].apply(extract_primary_topic)
    df["topic_confidence"] = df["topics"].apply(extract_topic_confidence)

    # Dry run
    if args.dry_run:
        print_dry_run(df, epoch, stats)
        return

    # Initialize matplotlib
    setup_matplotlib_defaults()

    # Build document
    print(f"Generating report: {output_path}")
    doc = Document()

    # Global document theme/layout
    configure_document_styles(doc)
    first_section = doc.sections[0]
    configure_section_layout(first_section, landscape=False)
    epoch_id = epoch.get("id", "?")
    set_document_header_footer(first_section, date_label, epoch_id)

    # --- Title Page ---
    build_title_page(doc, date_label, epoch)
    doc.add_page_break()

    # --- Table of Contents ---
    add_table_of_contents(doc)
    doc.add_page_break()

    # --- Executive Summary ---
    build_executive_summary(doc, df, epoch, stats)

    # --- Methodology ---
    build_methodology_block(doc, epoch, stats, date_label)

    # --- Topic Weights (if available) ---
    topic_weights = epoch.get("topic_weights")
    if topic_weights and isinstance(topic_weights, dict):
        doc.add_page_break()
        build_topic_weights_page(doc, epoch)

    # --- Topic Distribution ---
    doc.add_page_break()
    build_topic_page(doc, df)

    # --- Score Composition (landscape) ---
    add_landscape_section(doc, date_label, epoch_id)
    build_score_page(doc, df)

    # --- Classification Methods (portrait) ---
    restore_portrait_section(doc, date_label, epoch_id)
    build_classification_page(doc, df)

    # --- Author Diversity ---
    doc.add_page_break()
    build_diversity_page(doc, df)

    # --- Engagement Profile ---
    doc.add_page_break()
    build_engagement_page(doc, df)

    # --- Appendix (landscape) ---
    add_landscape_section(doc, date_label, epoch_id)
    build_sample_posts_page(
        doc,
        df.head(10),
        "Appendix A \u2014 Top 10 Sample Posts",
        "These examples are included for qualitative review and traceability. "
        "They are not exhaustive of all scored content.",
        is_appendix=True,
    )
    doc.add_page_break()
    build_sample_posts_page(
        doc,
        df.tail(10),
        "Appendix B \u2014 Bottom 10 Sample Posts",
        "Lower-ranked examples from the same top-1,000 analysis scope.",
        is_appendix=True,
    )

    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)

    file_size = os.path.getsize(output_path) / 1024
    print(f"Report saved: {output_path} ({file_size:.0f} KB)")


if __name__ == "__main__":
    main()
